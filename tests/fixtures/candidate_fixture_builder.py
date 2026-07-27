#!/usr/bin/env python3
"""Build and structurally verify synthetic resume golden-path fixtures.

Run only with an already trusted document runtime. The repository does not
install or declare a parser dependency from this fixture.
"""

from __future__ import annotations

import hashlib
import json
import re
import sys
import zipfile
from datetime import datetime
from pathlib import Path

import pdfplumber
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "resume_candidate_source.docx"
TAILORED = HERE / "resume_candidate_tailored.docx"
SOURCE_PDF = HERE / "resume_candidate_source.pdf"
TAILORED_PDF = HERE / "resume_candidate_tailored.pdf"
WRONG_CONTACT = HERE / "resume_candidate_wrong_contact.docx"
OVERFLOW = HERE / "resume_candidate_overflow.docx"
TWO_COLUMN = HERE / "resume_candidate_two_column.docx"
SCAN_PNG = HERE / "resume_candidate_scan_page.png"
SCAN_PDF = HERE / "resume_candidate_scanned.pdf"

NAME = "Rowan Example"
EMAIL = "rowan.candidate@example.invalid"
PHONE = "+1 202-555-0199"
WRONG_EMAIL = "other.candidate@example.invalid"

NAVY = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
FIXED_DOCUMENT_TIME = datetime(2026, 7, 27, 0, 0, 0)
FIXED_ZIP_TIME = (2026, 7, 27, 0, 0, 0)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def save_docx(doc: Document, path: Path):
    properties = doc.core_properties
    properties.author = "FindJobs synthetic fixture"
    properties.last_modified_by = "FindJobs synthetic fixture"
    properties.created = FIXED_DOCUMENT_TIME
    properties.modified = FIXED_DOCUMENT_TIME
    doc.save(path)

    # python-docx writes current ZIP member timestamps. Normalize the package so
    # repeated fixture builds produce identical bytes and stable content hashes.
    temporary = path.with_name(f".{path.name}.normalized")
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
            temporary, "w", compression=zipfile.ZIP_DEFLATED
        ) as target:
            for source_info in sorted(source.infolist(), key=lambda item: item.filename):
                normalized = zipfile.ZipInfo(source_info.filename, FIXED_ZIP_TIME)
                normalized.compress_type = zipfile.ZIP_DEFLATED
                normalized.create_system = source_info.create_system
                normalized.external_attr = source_info.external_attr
                target.writestr(normalized, source.read(source_info.filename))
        temporary.replace(path)
    finally:
        temporary.unlink(missing_ok=True)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, *, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_bullet_numbering(doc: Document, num_id=42):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(num_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_id = OxmlElement("w:abstractNumId")
    abstract_id.set(qn("w:val"), str(num_id))
    num.append(abstract_id)
    numbering.append(num)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after in (
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 14, 7),
        ("Heading 3", 12, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = NAVY
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    add_bullet_numbering(doc)


def add_identity(doc: Document, email=EMAIL):
    # Named resume_identity_block override of customer_pack: compact left title.
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    set_run(paragraph.add_run(NAME), size=22, bold=True, color=NAVY)
    contact = doc.add_paragraph()
    contact.paragraph_format.space_after = Pt(10)
    set_run(contact.add_run(f"{email} | {PHONE} | Shanghai, China"), size=10, color=GRAY)


def add_heading(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    set_run(paragraph.add_run(text), size=13, bold=True, color=NAVY)


def add_line(doc: Document, text: str, *, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    if bold_prefix and text.startswith(bold_prefix):
        set_run(paragraph.add_run(bold_prefix), bold=True)
        set_run(paragraph.add_run(text[len(bold_prefix):]))
    else:
        set_run(paragraph.add_run(text))
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "42")
    num_pr.extend([ilvl, num_id])
    paragraph.paragraph_format.keep_together = True
    set_run(paragraph.add_run(text))
    return paragraph


def add_source_content(doc: Document, email=EMAIL):
    add_identity(doc, email)
    add_heading(doc, "Summary")
    add_line(doc, "Statistics undergraduate with academic data-analysis and local LLM evaluation project experience.")

    add_heading(doc, "Education")
    add_line(doc, "Northbridge Institute of Technology (fictional) | B.Sc. Statistics | Expected June 2027")

    add_heading(doc, "Experience")
    add_line(doc, "Fictional Market Lab | Data Analysis Intern | May 2025 - August 2025")
    add_bullet(doc, "Worked from May 2025 through September 2025 on weekly marketplace analysis; this conflicts with the heading end date.")
    add_bullet(doc, "Analyzed 40,000 synthetic marketplace order records with SQL and Python to identify anomalous refund patterns.")
    add_bullet(doc, "Presented weekly findings to two research mentors and incorporated review notes.")

    add_heading(doc, "Projects")
    add_line(doc, "Storefront Risk Signals | Academic project using synthetic data")
    add_bullet(doc, "Built a logistic-regression prototype for a risk-screening exercise and documented false-positive tradeoffs.")
    add_line(doc, "LLM Support Triage | Academic project using synthetic questions")
    add_bullet(doc, "Built a local retrieval and evaluation prototype and compared answer-grounding failures.")

    add_heading(doc, "Skills and availability")
    add_line(doc, "Tools: SQL, Python, spreadsheets")
    add_line(doc, "English: professional working proficiency (self-reported)")
    add_line(doc, "Availability: September 2026 - December 2026 (four consecutive months)")


def build_source():
    doc = Document()
    configure_document(doc)
    add_source_content(doc)
    save_docx(doc, SOURCE)


def build_tailored(path: Path, email=EMAIL, overflow=False):
    doc = Document()
    configure_document(doc)
    add_identity(doc, email)
    add_heading(doc, "Profile")
    add_line(doc, "Statistics undergraduate with academic experience analyzing synthetic marketplace risk signals and building local LLM evaluation prototypes. Uses SQL and Python to investigate anomalies and communicate findings.")

    add_heading(doc, "Relevant projects")
    add_line(doc, "Storefront Risk Signals | Academic project using synthetic data")
    add_bullet(doc, "Used SQL and Python to analyze 40,000 synthetic marketplace orders and identify anomalous refund patterns.")
    add_bullet(doc, "Built a logistic-regression risk-screening prototype and documented false-positive tradeoffs; no production deployment is claimed.")
    add_line(doc, "LLM Support Triage | Academic project using synthetic questions")
    add_bullet(doc, "Built a local retrieval and evaluation prototype and compared answer-grounding failures.")

    add_heading(doc, "Experience")
    add_line(doc, "Fictional Market Lab | Data Analysis Intern | May 2025 - August 2025")
    add_bullet(doc, "Presented weekly marketplace findings to two research mentors and incorporated review notes.")

    add_heading(doc, "Education")
    add_line(doc, "Northbridge Institute of Technology (fictional) | B.Sc. Statistics | Expected June 2027")

    add_heading(doc, "Skills and availability")
    add_line(doc, "Tools: SQL, Python, spreadsheets")
    add_line(doc, "English: professional working proficiency (self-reported)")
    add_line(doc, "Availability: September 2026 - December 2026 (four consecutive months)")

    if overflow:
        add_heading(doc, "Deliberate overflow fixture - must fail upload gate")
        for index in range(1, 75):
            add_bullet(doc, f"Overflow probe {index}: repeated synthetic text used only to force unexpected page-count growth during QA.")
    save_docx(doc, path)


def build_two_column():
    doc = Document()
    configure_document(doc)
    section = doc.sections[0]
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")
    add_identity(doc)
    add_heading(doc, "Deliberate two-column ambiguity")
    for index in range(1, 17):
        add_line(doc, f"Column-flow probe {index}: synthetic content for reading-order detection.")
    save_docx(doc, TWO_COLUMN)


def build_text_pdf(path: Path, *, tailored: bool):
    pdf = canvas.Canvas(
        str(path), pagesize=letter, invariant=1, pageCompression=1
    )
    text = pdf.beginText(72, 730)
    text.setFont("Helvetica", 10)
    text.setLeading(15)
    if tailored:
        lines = [
            NAME,
            f"{EMAIL} | {PHONE} | Shanghai, China",
            "PROFILE",
            "Statistics undergraduate with academic experience analyzing synthetic marketplace risk signals.",
            "RELEVANT PROJECTS",
            "Storefront Risk Signals | Academic project using synthetic data",
            "Used SQL and Python to analyze 40,000 synthetic marketplace orders and identify anomalous refunds.",
            "Built a logistic-regression risk-screening prototype; no production deployment is claimed.",
            "LLM Support Triage | Academic project using synthetic questions",
            "Built a local retrieval and evaluation prototype and compared answer-grounding failures.",
            "EXPERIENCE",
            "Fictional Market Lab | Data Analysis Intern | May 2025 - August 2025",
            "Presented weekly findings to two research mentors and incorporated review notes.",
            "EDUCATION",
            "Northbridge Institute of Technology (fictional) | B.Sc. Statistics | Expected June 2027",
            "SKILLS AND AVAILABILITY",
            "Tools: SQL, Python, spreadsheets",
            "English: professional working proficiency (self-reported)",
            "Availability: September 2026 - December 2026 (four consecutive months)",
        ]
    else:
        lines = [
            NAME,
            f"{EMAIL} | {PHONE} | Shanghai, China",
            "SUMMARY",
            "Statistics undergraduate with academic data-analysis and local LLM evaluation experience.",
            "EDUCATION",
            "Northbridge Institute of Technology (fictional) | B.Sc. Statistics | Expected June 2027",
            "EXPERIENCE",
            "Fictional Market Lab | Data Analysis Intern | May 2025 - August 2025",
            "Worked from May 2025 through September 2025; this conflicts with the heading end date.",
            "Analyzed 40,000 synthetic marketplace order records with SQL and Python.",
            "Presented weekly findings to two research mentors and incorporated review notes.",
            "PROJECTS",
            "Storefront Risk Signals | Academic project using synthetic data",
            "Built a logistic-regression risk-screening prototype and documented false-positive tradeoffs.",
            "LLM Support Triage | Academic project using synthetic questions",
            "Built a local retrieval and evaluation prototype and compared answer-grounding failures.",
            "SKILLS AND AVAILABILITY",
            "Tools: SQL, Python, spreadsheets",
            "English: professional working proficiency (self-reported)",
            "Availability: September 2026 - December 2026 (four consecutive months)",
        ]
    for line in lines:
        text.textLine(line)
    pdf.drawText(text)
    pdf.showPage()
    pdf.save()


def build_scanned_pdf():
    width, height = 1700, 2200
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default(size=32)
    lines = [
        "ROWAN EXAMPLE - IMAGE-ONLY SCAN FIXTURE",
        EMAIL,
        PHONE,
        "This page intentionally has no PDF text layer.",
        "Expected result: needs_ocr; do not match or tailor.",
    ]
    y = 180
    for line in lines:
        draw.text((150, y), line, fill="black", font=font)
        y += 85
    image.save(SCAN_PNG)
    pdf = canvas.Canvas(str(SCAN_PDF), pagesize=letter, invariant=1, pageCompression=1)
    pdf.drawImage(ImageReader(image), 0, 0, width=letter[0], height=letter[1])
    pdf.showPage()
    pdf.save()


def extract_docx(path: Path) -> str:
    doc = Document(path)
    chunks = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def extract_pdf(path: Path) -> tuple[str, int]:
    with pdfplumber.open(path) as pdf:
        text = "\n".join((page.extract_text() or "") for page in pdf.pages).strip()
        return text, len(pdf.pages)


def detect_columns(path: Path) -> int:
    doc = Document(path)
    counts = []
    for section in doc.sections:
        cols = section._sectPr.xpath("./w:cols")[0]
        counts.append(int(cols.get(qn("w:num"), "1")))
    return max(counts)


def write_contracts():
    source_hash = sha256(SOURCE)
    tailored_hash = sha256(TAILORED)
    url = "https://jobs.bytedance.com/campus/position/7665368414857726261/detail"
    captured = "2026-07-27T15:15:58Z"

    job = {
        "schemaVersion": "1.0",
        "jobId": "job-bytedance-A106199",
        "canonicalKey": url,
        "platformJobId": "A106199",
        "aliases": [url],
        "source": {"site": "ByteDance Careers", "url": url, "capturedAt": captured, "agentId": "golden-path-read-only"},
        "company": "ByteDance",
        "title": "Risk Strategy Data Analysis Intern - TikTok Shop",
        "location": "Shanghai",
        "employmentType": "internship",
        "recruitmentType": "internship",
        "graduationCohorts": [],
        "requirements": {
            "must": [
                "Current bachelor-or-higher student and available for at least three months",
                "Risk-strategy data analysis, monitoring indicators, and alerts",
                "Risk profiling and model or strategy development",
                "Special analysis and strategy iteration recommendations",
                "Independent model design/development and implementation",
                "Cross-functional collaboration",
                "Data analysis that finds insights, risk, and anomalies in complex data",
                "AI/LLM application practice",
                "Communication, collaboration, rigor, and English as a working language",
            ],
            "preferred": ["Mathematics, statistics, or computer-science field", "Risk-control experience"],
            "education": ["Bachelor degree or above, currently enrolled"],
            "experience": ["Internship availability of at least three months"],
            "salary": None,
            "deadline": None,
        },
        "descriptionExcerpt": "International e-commerce merchant risk strategy data analysis; risk profiles/models; special analysis; independent model development and implementation; cross-functional collaboration.",
        "evidenceRefs": [
            {"id": "JD-1", "url": url, "capturedAt": captured, "excerpt": "Daily internship for current students, three months or more."},
            {"id": "JD-2", "url": url, "capturedAt": captured, "excerpt": "Analyze merchant risk strategy data; evaluate risk; design monitoring indicators and alerts."},
            {"id": "JD-3", "url": url, "capturedAt": captured, "excerpt": "Mine merchant risk profiles; build models or strategies; provide analysis and iteration recommendations."},
            {"id": "JD-4", "url": url, "capturedAt": captured, "excerpt": "Independently design/develop models and drive implementation; collaborate with product, operations, algorithm, and engineering."},
            {"id": "JD-5", "url": url, "capturedAt": captured, "excerpt": "Current bachelor-or-higher student; math/statistics/computer science preferred."},
            {"id": "JD-6", "url": url, "capturedAt": captured, "excerpt": "Data analysis, complex-data insight/risk/anomaly detection, AI/LLM practice, communication, rigor, working English; risk-control experience preferred."},
        ],
        "deduplication": {"decision": "unique", "primaryJobId": None, "evidenceRefs": ["JD-1"]},
        "freshness": "current",
        "status": "candidate",
    }

    evidence = [
        {"id": "RES-1", "location": "resume-original#paragraph:1", "exactFact": NAME, "confidence": "high"},
        {"id": "RES-2", "location": "resume-original#paragraph:2", "exactFact": f"{EMAIL} | {PHONE} | Shanghai, China", "confidence": "high"},
        {"id": "RES-3", "location": "resume-original#paragraph:6", "exactFact": "B.Sc. Statistics | Expected June 2027", "confidence": "high"},
        {"id": "RES-4", "location": "resume-original#paragraph:8", "exactFact": "Data Analysis Intern | May 2025 - August 2025", "confidence": "high"},
        {"id": "RES-5", "location": "resume-original#paragraph:9", "exactFact": "Worked from May 2025 through September 2025", "confidence": "high"},
        {"id": "RES-6", "location": "resume-original#paragraph:10", "exactFact": "Analyzed 40,000 synthetic marketplace order records with SQL and Python to identify anomalous refund patterns.", "confidence": "high"},
        {"id": "RES-7", "location": "resume-original#paragraph:11", "exactFact": "Presented weekly findings to two research mentors and incorporated review notes.", "confidence": "high"},
        {"id": "RES-8", "location": "resume-original#paragraph:14", "exactFact": "Built a logistic-regression prototype for a risk-screening exercise and documented false-positive tradeoffs.", "confidence": "high"},
        {"id": "RES-9", "location": "resume-original#paragraph:16", "exactFact": "Built a local retrieval and evaluation prototype and compared answer-grounding failures.", "confidence": "high"},
        {"id": "RES-10", "location": "resume-original#paragraph:18", "exactFact": "Tools: SQL, Python, spreadsheets", "confidence": "high"},
        {"id": "RES-11", "location": "resume-original#paragraph:19", "exactFact": "English: professional working proficiency (self-reported)", "confidence": "medium"},
        {"id": "RES-12", "location": "resume-original#paragraph:20", "exactFact": "Availability: September 2026 - December 2026 (four consecutive months)", "confidence": "high"},
    ]
    profile = {
        "schemaVersion": "1.0",
        "profileId": "candidate-rowan-example",
        "sourceDocuments": [{"id": "resume-original", "path": "tests/fixtures/resume_candidate_source.docx", "sha256": source_hash}],
        "identity": {"name": NAME, "contact": {"email": EMAIL, "phone": PHONE}},
        "preferences": {"roles": ["data analysis internship"], "locations": ["Shanghai"], "workModes": [], "recruitmentPrograms": ["internship"], "graduationCohorts": ["2027"], "salary": None, "startDate": "September 2026"},
        "fieldScopes": {"salary": "filter_only", "locations": "filter_only", "workAuthorization": "filter_only"},
        "education": [{"school": "Northbridge Institute of Technology (fictional)", "degree": "B.Sc. Statistics", "graduation": "Expected June 2027", "provenance": "observed", "evidenceRefs": ["RES-3"]}],
        "experience": [{"organization": "Fictional Market Lab", "title": "Data Analysis Intern", "dates": ["May 2025 - August 2025", "May 2025 - September 2025"], "provenance": "observed", "evidenceRefs": ["RES-4", "RES-5", "RES-6", "RES-7"]}],
        "projects": [
            {"name": "Storefront Risk Signals", "experienceType": "academic_synthetic", "provenance": "observed", "evidenceRefs": ["RES-6", "RES-8"]},
            {"name": "LLM Support Triage", "experienceType": "academic_synthetic", "provenance": "observed", "evidenceRefs": ["RES-9"]},
        ],
        "skills": {
            "technical": [{"name": "SQL", "level": "demonstrated", "evidenceRefs": ["RES-6", "RES-10"]}, {"name": "Python", "level": "demonstrated", "evidenceRefs": ["RES-6", "RES-10"]}],
            "domain": [{"name": "risk screening", "level": "academic_synthetic", "evidenceRefs": ["RES-8"]}],
            "tools": [{"name": "spreadsheets", "level": "claimed", "evidenceRefs": ["RES-10"]}],
            "languages": [{"name": "English", "level": "claimed_only", "evidenceRefs": ["RES-11"]}],
        },
        "constraints": [{"kind": "availability", "value": "September 2026 - December 2026", "provenance": "observed", "evidenceRefs": ["RES-12"]}],
        "evidenceRefs": evidence,
        "missingFacts": ["work authorization", "salary expectation"],
        "conflicts": [{"field": "experience[0].endDate", "values": ["August 2025", "September 2025"], "evidenceRefs": ["RES-4", "RES-5"], "affectsEligibility": False}],
        "extractionDiagnostics": {"status": "ready", "pageCount": 1, "renderedPages": ["page-1"], "textLayerByPage": [True], "ambiguousLocations": [], "externalRelationshipsIgnored": [], "injectionFlags": []},
    }

    assessments = [
        ("M1", "Risk strategy analysis plus monitoring indicators/alerts", "partial", 0.5, ["JD-2", "RES-6"]),
        ("M2", "Risk profiling and model/strategy development", "partial", 0.5, ["JD-3", "RES-8"]),
        ("M3", "Special analysis and strategy iteration recommendations", "partial", 0.5, ["JD-3", "RES-8"]),
        ("M4", "Independent model development and implementation", "partial", 0.5, ["JD-4", "RES-8"]),
        ("M5", "Cross-functional collaboration", "partial", 0.5, ["JD-4", "RES-7"]),
        ("M6", "Complex-data insight, risk, and anomaly analysis", "demonstrated", 1.0, ["JD-6", "RES-6"]),
        ("M7", "AI/LLM application practice", "demonstrated", 1.0, ["JD-6", "RES-9"]),
        ("M8", "Communication and rigor", "partial", 0.5, ["JD-6", "RES-7"]),
        ("M9", "English as a working language", "claimed_only", 0.5, ["JD-6", "RES-11"]),
        ("P1", "Statistics-related field", "demonstrated", 1.0, ["JD-5", "RES-3"]),
        ("P2", "Risk-control experience", "partial", 0.5, ["JD-6", "RES-8"]),
    ]
    normalized = []
    for req_id, requirement, status, multiplier, refs in assessments:
        is_must = req_id.startswith("M")
        weight = 70 / 9 if is_must else 30 / 2
        normalized.append({"id": req_id, "classification": "must" if is_must else "preferred", "requirement": requirement, "status": status, "weight": round(weight, 4), "multiplier": multiplier, "evidenceRefs": refs})

    match = {
        "schemaVersion": "1.0",
        "profileId": "candidate-rowan-example",
        "jobId": "job-bytedance-A106199",
        "decision": "possible_match",
        "score": {"value": 65, "scale": "0-100", "method": "9 must requirements share 70 points; 2 preferred share 30; status multipliers applied before final rounding", "calculation": "(5.5/9*70) + (1.5/2*30) = 65.277... -> 65"},
        "hardConstraints": [
            {"criterion": "Current bachelor-or-higher student", "status": "met", "evidenceRefs": ["JD-1", "JD-5", "RES-3"]},
            {"criterion": "Available at least three months", "status": "met", "evidenceRefs": ["JD-1", "RES-12"]},
        ],
        "normalizedRequirements": normalized,
        "strengths": [
            {"claim": "Statistics education and demonstrated synthetic marketplace anomaly analysis", "evidenceRefs": ["JD-5", "JD-6", "RES-3", "RES-6"]},
            {"claim": "Documented AI/LLM application project", "evidenceRefs": ["JD-6", "RES-9"]},
        ],
        "gaps": [
            {"claim": "No evidence of production deployment or independently driving model implementation", "severity": "material", "evidenceRefs": ["JD-4", "RES-8"]},
            {"claim": "Risk-control evidence is academic and synthetic, not commercial", "severity": "minor", "evidenceRefs": ["JD-6", "RES-8"]},
            {"claim": "English level is self-reported without demonstrated work artifact", "severity": "minor", "evidenceRefs": ["JD-6", "RES-11"]},
        ],
        "confidence": "medium",
        "confidenceRationale": "Hard constraints have direct evidence; multiple must requirements are only partially demonstrated and one language item is claimed-only.",
        "injectionFlags": [],
        "userQuestions": [],
        "generatedAt": "2026-07-27T15:30:00Z",
    }

    user_answers = {
        "schemaVersion": "1.0",
        "profileId": "candidate-rowan-example",
        "answers": [{"evidenceId": "USER-A1", "question": "Which internship end date is correct?", "answer": "August 2025", "provenance": "user_provided", "syntheticFixture": True}],
    }

    changes = [
        {"before": "Statistics undergraduate with academic data-analysis and local LLM evaluation project experience.", "after": "Statistics undergraduate with academic experience analyzing synthetic marketplace risk signals and building local LLM evaluation prototypes. Uses SQL and Python to investigate anomalies and communicate findings.", "evidenceRefs": ["RES-3", "RES-6", "RES-7", "RES-9"]},
        {"before": "Analyzed 40,000 synthetic marketplace order records with SQL and Python to identify anomalous refund patterns.", "after": "Used SQL and Python to analyze 40,000 synthetic marketplace orders and identify anomalous refund patterns.", "evidenceRefs": ["RES-6"]},
        {"before": "Built a logistic-regression prototype for a risk-screening exercise and documented false-positive tradeoffs.", "after": "Built a logistic-regression risk-screening prototype and documented false-positive tradeoffs; no production deployment is claimed.", "evidenceRefs": ["RES-8"]},
        {"before": "May 2025 - August/September 2025 conflict", "after": "May 2025 - August 2025", "evidenceRefs": ["RES-4", "RES-5", "USER-A1"]},
    ]
    draft = {
        "schemaVersion": "1.0",
        "draftId": "draft-bytedance-A106199-rowan-example",
        "jobId": "job-bytedance-A106199",
        "profileId": "candidate-rowan-example",
        "resumeVersion": f"resume-{tailored_hash}",
        "files": [{"path": "tests/fixtures/resume_candidate_tailored.docx", "kind": "resume", "sha256": tailored_hash}],
        "answers": [],
        "changes": changes,
        "riskFlags": ["Commercial risk-control experience is not claimed", "Production model deployment is not claimed"],
        "missingRequiredFields": ["work authorization", "salary expectation"],
        "readyForReview": False,
        "claimLedger": [
            {"claimId": "CL-1", "exactFact": "40,000 synthetic marketplace orders analyzed with SQL and Python for anomalous refund patterns", "allowedTransformations": ["shorten", "active voice", "reorder"], "candidateEvidence": ["RES-6"], "risk": "Must remain synthetic"},
            {"claimId": "CL-2", "exactFact": "Academic logistic-regression risk-screening prototype; false-positive tradeoffs documented", "allowedTransformations": ["shorten", "reorder"], "candidateEvidence": ["RES-8"], "risk": "No production deployment"},
            {"claimId": "CL-3", "exactFact": "Local LLM retrieval/evaluation prototype on synthetic questions", "allowedTransformations": ["shorten", "reorder"], "candidateEvidence": ["RES-9"], "risk": "No commercial deployment"},
        ],
        "qa": {"sourceSha256": source_hash, "outputSha256": tailored_hash, "readbackPassed": True, "criticalValuesPassed": True, "contactIsolationPassed": True, "claimReconciliationPassed": True, "renderedPages": ["page-1"], "layoutPassed": True, "issues": []},
    }

    mutation = {
        "syntheticMurphyCase": "fabricated_metric",
        "proposedText": "Reduced merchant fraud losses by 31% in production.",
        "expected": "reject",
        "reason": "31%, fraud-loss impact, merchant scope, and production deployment have no candidate evidence",
    }

    files = {
        "candidate_job_posting.json": job,
        "candidate_profile.json": profile,
        "candidate_match_report.json": match,
        "candidate_user_answers.json": user_answers,
        "candidate_application_draft.json": draft,
        "candidate_claim_mutation.json": mutation,
    }
    for filename, payload in files.items():
        (HERE / filename).write_text(json.dumps(payload, indent=2, ensure_ascii=True) + "\n", encoding="utf-8")


def structural_checks():
    source_text = extract_docx(SOURCE)
    tailored_text = extract_docx(TAILORED)
    wrong_text = extract_docx(WRONG_CONTACT)
    source_pdf_text, source_pdf_pages = extract_pdf(SOURCE_PDF)
    tailored_pdf_text, tailored_pdf_pages = extract_pdf(TAILORED_PDF)
    scan_text, scan_pages = extract_pdf(SCAN_PDF)

    contacts = set(re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+", tailored_text))
    wrong_contacts = set(re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+", wrong_text))
    results = {
        "source_sha256": sha256(SOURCE),
        "tailored_sha256": sha256(TAILORED),
        "source_readback_contains_metric": "40,000 synthetic marketplace order records" in source_text,
        "source_readback_contains_conflicting_dates": "August 2025" in source_text and "September 2025" in source_text,
        "tailored_contact_exact": contacts == {EMAIL},
        "wrong_contact_fixture_rejected": wrong_contacts != {EMAIL},
        "source_pdf_text_layer_readback": source_pdf_pages == 1 and "40,000 synthetic marketplace order records" in source_pdf_text,
        "tailored_pdf_text_layer_readback": tailored_pdf_pages == 1 and "40,000 synthetic marketplace orders" in tailored_pdf_text,
        "tailored_pdf_contact_exact": set(re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+", tailored_pdf_text)) == {EMAIL},
        "two_column_detected": detect_columns(TWO_COLUMN) == 2,
        "scanned_pdf_has_no_text_layer": scan_pages == 1 and scan_text == "",
        "fabricated_metric_absent_from_tailored": "31%" not in tailored_text and "fraud losses" not in tailored_text,
    }
    results["all_structural_checks_passed"] = all(value for key, value in results.items() if key not in {"source_sha256", "tailored_sha256"})
    (HERE / "candidate_structural_results.json").write_text(json.dumps(results, indent=2) + "\n", encoding="utf-8")
    return results


def main():
    if len(sys.argv) > 1 and sys.argv[1] not in {"build", "verify"}:
        raise SystemExit("usage: candidate_fixture_builder.py [build|verify]")
    if len(sys.argv) == 1 or sys.argv[1] == "build":
        build_source()
        build_tailored(TAILORED)
        build_text_pdf(SOURCE_PDF, tailored=False)
        build_text_pdf(TAILORED_PDF, tailored=True)
        build_tailored(WRONG_CONTACT, email=WRONG_EMAIL)
        build_tailored(OVERFLOW, overflow=True)
        build_two_column()
        build_scanned_pdf()
        write_contracts()
    results = structural_checks()
    print(json.dumps(results, indent=2))
    if not results["all_structural_checks_passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
